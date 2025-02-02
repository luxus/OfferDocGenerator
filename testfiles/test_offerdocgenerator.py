import sys
import os
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
import unittest
import tempfile
import shutil
from pathlib import Path
import yaml
import docx
from unittest import mock
import offerdocgenerator

class TestOfferDocGenerator(unittest.TestCase):
    def setUp(self):
        """Set up test fixtures."""
        # Set up test directories
        self.test_dir = Path(__file__).parent
        self.config_file = self.test_dir / "test_config.yaml"
        self.templates_dir = self.test_dir / "templates"
        self.output_dir = self.test_dir / "output"
        self.textblocks_dir = self.test_dir / "textblocks"
        
        # Clean previous test files but preserve output
        shutil.rmtree(self.templates_dir, ignore_errors=True)
        shutil.rmtree(self.textblocks_dir, ignore_errors=True)
        
        # Create output dir if needed but don't clean it
        self.output_dir.mkdir(exist_ok=True)
        self.product_name = "Web Application Security Assessment"

        # Create necessary directories
        self.templates_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)
        (self.textblocks_dir / "common").mkdir(parents=True, exist_ok=True)
        (self.textblocks_dir / "products" / self.product_name).mkdir(parents=True, exist_ok=True)

        # Create common textblocks
        self._create_textblock_file(
            self.textblocks_dir / "common" / "section_1_1_EN.docx",
            "Our standard security assessment provides a comprehensive evaluation of your web application's security posture."
        )
        self._create_textblock_file(
            self.textblocks_dir / "common" / "section_1_1_DE.docx",
            "Unsere Standard-Sicherheitsbewertung bietet eine umfassende Evaluation der Sicherheitslage Ihrer Webanwendung."
        )

        # Create product-specific textblocks for first product
        self._create_textblock_file(
            self.textblocks_dir / "products" / self.product_name / "section_1_1_1_EN.docx",
            """The Web Application Security Assessment includes:

- Vulnerability scanning
- Manual penetration testing
- Code review"""
        )
        self._create_textblock_file(
            self.textblocks_dir / "products" / self.product_name / "section_1_1_1_DE.docx",
            """Die Web Application Security Assessment beinhaltet:

- Schwachstellenscanning
- Manuelle Penetrationstests
- Code-Review"""
        )

        # Create second product and its textblocks
        self.product_name2 = "Network Security Audit"
        (self.textblocks_dir / "products" / self.product_name2).mkdir(parents=True, exist_ok=True)
        self._create_textblock_file(
            self.textblocks_dir / "products" / self.product_name2 / "section_1_1_1_EN.docx",
            """The Network Security Audit includes:

- Firewall configuration review
- Intrusion detection system testing
- Security policy evaluation"""
        )
        self._create_textblock_file(
            self.textblocks_dir / "products" / self.product_name2 / "section_1_1_1_DE.docx",
            """Die Netzwerksicherheitsprüfung umfasst:

- Überprüfung der Firewall-Konfiguration
- Test des Intrusion Detection Systems
- Bewertung der Sicherheitsrichtlinien"""
        )

        # Create base templates for EN and DE
        # English template
        self.template_file_en = self.templates_dir / "base_EN.docx"
        doc = docx.Document()
        doc.add_heading('Offer: {{ Offer.number }}', 0)
        doc.add_paragraph('Date: {{ Offer.date }}')
        doc.add_paragraph('Valid for: {{ Offer.validity }}')
        doc.add_heading('Customer Information', 1)
        doc.add_paragraph('{{ Customer.name }}')
        doc.add_paragraph('{{ Customer.address }}')
        doc.add_paragraph('{{ Customer.city }}, {{ Customer.zip }}')
        doc.add_paragraph('{{ Customer.country }}')
        doc.add_heading('Product Description', 1)
        p = doc.add_paragraph()
        p.add_run('{{r section_1_1 }}')
        doc.add_heading('Detailed Scope', 2)
        p = doc.add_paragraph()
        p.add_run('{{r section_1_1_1 }}')
        doc.add_paragraph('Total Price: {{ Offer.currency }} 10,000')
        doc.add_heading('Sales Contact', 1)
        doc.add_paragraph('{{ Sales.name }}')
        doc.add_paragraph('{{ Sales.email }}')
        doc.add_paragraph('{{ Sales.phone }}')
        doc.save(str(self.template_file_en))

        # German template
        self.template_file_de = self.templates_dir / "base_de.docx"
        doc = docx.Document()
        doc.add_heading('Angebot: {{ Offer.number }}', 0)
        doc.add_paragraph('Datum: {{ Offer.date }}')
        doc.add_paragraph('Gültig für: {{ Offer.validity }}')
        doc.add_heading('Kundeninformationen', 1)
        doc.add_paragraph('{{ Customer.name }}')
        doc.add_paragraph('{{ Customer.address }}')
        doc.add_paragraph('{{ Customer.city }}, {{ Customer.zip }}')
        doc.add_paragraph('{{ Customer.country }}')
        doc.add_heading('Produktbeschreibung', 1)
        p = doc.add_paragraph()
        p.add_run('{{r section_1_1 }}')
        doc.add_heading('Detaillierter Umfang', 2)
        p = doc.add_paragraph()
        p.add_run('{{r section_1_1_1 }}')
        doc.add_paragraph('Gesamtpreis: {{ Offer.currency }} 10.000')
        doc.add_heading('Vertriebskontakt', 1)
        doc.add_paragraph('{{ Sales.name }}')
        doc.add_paragraph('{{ Sales.email }}')
        doc.add_paragraph('{{ Sales.phone }}')
        doc.save(str(self.template_file_de))

        # Create config file
        config = {
            "offer": {
                "sections": ["1_1", "1_1_1"],
                "template": str(self.templates_dir)
            },
            "textblocks": {
                "common": {
                    "folder": str(self.textblocks_dir / "common")
                },
                "products_dir": str(self.textblocks_dir / "products")
            },
            "output": {
                "folder": str(self.output_dir)
            }
        }
        with open(self.config_file, 'w') as f:
            yaml.dump(config, f)

    def _create_textblock_file(self, file_path: Path, content: str):
        """Helper method to create a docx file with given content."""
        doc = docx.Document()
        # Add each paragraph with proper styling
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                p = doc.add_paragraph()
                # If it's a bullet point, use a list style
                if paragraph.strip().startswith('-'):
                    p.style = 'List Bullet'
                    p.text = paragraph.strip()[2:]  # Remove the '- ' prefix
                else:
                    p.text = paragraph.strip()
        doc.save(str(file_path))

    def tearDown(self):
        """Keep generated output files, only clean up test inputs."""
        # Remove config file
        if self.config_file.exists():
            self.config_file.unlink()
        
        # Clean up test template and textblock files but keep output
        cleanup_files = [
            self.template_file_en,
            self.template_file_de,
            self.textblocks_dir / "common" / "section_1_1_EN.docx",
            self.textblocks_dir / "common" / "section_1_1_DE.docx",
            self.textblocks_dir / "products" / self.product_name / "section_1_1_1_EN.docx",
            self.textblocks_dir / "products" / self.product_name / "section_1_1_1_DE.docx",
            self.textblocks_dir / "products" / self.product_name2 / "section_1_1_1_EN.docx",
            self.textblocks_dir / "products" / self.product_name2 / "section_1_1_1_DE.docx"
        ]
        
        for file_path in cleanup_files:
            if file_path.exists():
                file_path.unlink()

        # Print location of preserved output files
        print(f"\nGenerated files preserved in: {self.output_dir}")

    def test_get_product_names(self):
        """Test that product names are correctly detected from the directory structure."""
        products = offerdocgenerator.get_product_names(Path(self.test_dir) / "textblocks")
        self.assertEqual(len(products), 2)
        self.assertIn(self.product_name, products)
        self.assertIn(self.product_name2, products)

    def test_load_config(self):
        """Test that the YAML configuration is loaded correctly."""
        config = offerdocgenerator.load_config(self.config_file)
        self.assertIsInstance(config, offerdocgenerator.Config)
        self.assertEqual(config.offer["sections"], ["1_1", "1_1_1"])
        self.assertIn("common", config.textblocks)
        self.assertEqual(config.textblocks["common"]["folder"], str(self.textblocks_dir / "common"))

    def test_load_textblocks(self):
        """Test loading of textblocks from both product-specific and common directories."""
        config = offerdocgenerator.load_config(self.config_file)
        sections = ["1_1", "1_1_1"]

        # Test German textblocks
        textblocks_de = offerdocgenerator.load_textblocks(config, sections, self.product_name, "DE")
        self.assertIn("section_1_1", textblocks_de)
        self.assertIn("section_1_1_1", textblocks_de)
        self.assertIn("Sicherheitsbewertung", str(textblocks_de["section_1_1"]))
        self.assertIn("Schwachstellenscanning", str(textblocks_de["section_1_1_1"]))

        # Test English textblocks
        textblocks_en = offerdocgenerator.load_textblocks(config, sections, self.product_name, "EN")
        self.assertIn("section_1_1", textblocks_en)
        self.assertIn("section_1_1_1", textblocks_en)
        self.assertIn("comprehensive evaluation", str(textblocks_en["section_1_1"]))
        self.assertIn("Vulnerability scanning", str(textblocks_en["section_1_1_1"]))

    def test_render_offer(self):
        """Test rendering for all language/currency combinations."""
        config = offerdocgenerator.load_config(self.config_file)
        products = offerdocgenerator.get_product_names(Path(self.test_dir) / "textblocks")
        
        # Verify files per product
        for product in products:
            for lang in ["EN", "DE"]:
                for currency in ["CHF", "EUR"]:
                    with self.subTest(product=product, language=lang, currency=currency):
                        # Build context with currency
                        context = offerdocgenerator.build_context(config, lang, product, currency)
                        textblocks = offerdocgenerator.load_textblocks(config, config.offer["sections"], product, lang)
                        context.update(textblocks)
                        
                        # Select template
                        template = self.template_file_en if lang == "EN" else self.template_file_de
                        output_file = self.output_dir / f"Offer_{product}_{lang}_{currency}.docx"
                        
                        # Render and verify
                        offerdocgenerator.render_offer(template, context, output_file)
                        self.assertTrue(output_file.exists())
                        
                        # Verify currency in document
                        doc = docx.Document(str(output_file))
                        full_text = "\n".join(para.text for para in doc.paragraphs)
                        self.assertIn(currency, full_text)
        
        # Verify total file count (2 products * 2 langs * 2 currencies = 8 files)
        generated_files = list(self.output_dir.glob("Offer_*.docx"))
        self.assertEqual(len(generated_files), 8,
                        f"Expected 8 files for 2 products, found {len(generated_files)}")

    def test_create_test_files(self):
        """Just create the test files and directories."""
        self.test_dir = Path(__file__).parent
        self.config_file = self.test_dir / "config.yaml"
        self.templates_dir = self.test_dir / "templates"
        self.output_dir = Path("/Users/luxus/projects/OfferDocGenerator/output")
        self.textblocks_dir = self.test_dir / "textblocks"
        self.product_name = "Web Application Security Assessment"

        # Clean up existing directories first
        for dir_path in [self.templates_dir, self.textblocks_dir]:
            if dir_path.exists():
                shutil.rmtree(dir_path)

        # Create necessary directories
        print("\nCreating directories...")
        self.templates_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)
        (self.textblocks_dir / "common").mkdir(parents=True, exist_ok=True)
        (self.textblocks_dir / "products" / self.product_name).mkdir(parents=True, exist_ok=True)

        # Create common textblocks
        print("\nCreating common textblocks...")
        common_en = self.textblocks_dir / "common" / "section_1_1_EN.docx"
        common_de = self.textblocks_dir / "common" / "section_1_1_DE.docx"
        
        self._create_textblock_file(
            common_en,
            "Our standard security assessment provides a comprehensive evaluation of your web application's security posture."
        )
        self._create_textblock_file(
            common_de,
            "Unsere Standard-Sicherheitsbewertung bietet eine umfassende Evaluation der Sicherheitslage Ihrer Webanwendung."
        )

        # Create product-specific textblocks
        print("\nCreating product-specific textblocks...")
        product_en = self.textblocks_dir / "products" / self.product_name / "section_1_1_1_EN.docx"
        product_de = self.textblocks_dir / "products" / self.product_name / "section_1_1_1_DE.docx"
        
        self._create_textblock_file(
            product_en,
            """The Web Application Security Assessment includes:

- Vulnerability scanning
- Manual penetration testing
- Code review"""
        )
        self._create_textblock_file(
            product_de,
            """Die Web Application Security Assessment beinhaltet:

- Schwachstellenscanning
- Manuelle Penetrationstests
- Code-Review"""
        )

        # Create base templates for EN and DE
        print("\nCreating templates...")
        # English template
        self.template_file_en = self.templates_dir / "base_EN.docx"
        doc = docx.Document()
        doc.add_heading('Offer: {{ Offer.number }}', 0)
        doc.add_paragraph('Date: {{ Offer.date }}')
        doc.add_paragraph('Valid for: {{ Offer.validity }}')
        doc.add_heading('Customer Information', 1)
        doc.add_paragraph('{{ Customer.name }}')
        doc.add_paragraph('{{ Customer.address }}')
        doc.add_paragraph('{{ Customer.city }}, {{ Customer.zip }}')
        doc.add_paragraph('{{ Customer.country }}')
        doc.add_heading('Product Description', 1)
        p = doc.add_paragraph()
        p.add_run('{{ section_1_1 }}')
        doc.add_heading('Detailed Scope', 2)
        p = doc.add_paragraph()
        p.add_run('{{ section_1_1_1 }}')
        doc.add_heading('Sales Contact', 1)
        doc.add_paragraph('{{ Sales.name }}')
        doc.add_paragraph('{{ Sales.email }}')
        doc.add_paragraph('{{ Sales.phone }}')
        doc.save(str(self.template_file_en))

        # German template
        self.template_file_de = self.templates_dir / "base_DE.docx"
        doc = docx.Document()
        doc.add_heading('Angebot: {{ Offer.number }}', 0)
        doc.add_paragraph('Datum: {{ Offer.date }}')
        doc.add_paragraph('Gültig für: {{ Offer.validity }}')
        doc.add_heading('Kundeninformationen', 1)
        doc.add_paragraph('{{ Customer.name }}')
        doc.add_paragraph('{{ Customer.address }}')
        doc.add_paragraph('{{ Customer.city }}, {{ Customer.zip }}')
        doc.add_paragraph('{{ Customer.country }}')
        doc.add_heading('Produktbeschreibung', 1)
        p = doc.add_paragraph()
        p.add_run('{{ section_1_1 }}')
        doc.add_heading('Detaillierter Umfang', 2)
        p = doc.add_paragraph()
        p.add_run('{{ section_1_1_1 }}')
        doc.add_heading('Vertriebskontakt', 1)
        doc.add_paragraph('{{ Sales.name }}')
        doc.add_paragraph('{{ Sales.email }}')
        doc.add_paragraph('{{ Sales.phone }}')
        doc.save(str(self.template_file_de))

        # Create config file
        print("\nCreating config file...")
        config = {
            'offer': {
                'sections': ['1_1', '1_1_1'],
                'template': str(self.templates_dir)
            },
            'textblocks': {
                'common': {
                    'folder': str(self.textblocks_dir / "common")
                },
                'products_dir': str(self.textblocks_dir / "products")
            },
            'output': {
                'folder': str(self.output_dir)
            }
        }
        
        with open(self.config_file, 'w') as f:
            yaml.dump(config, f)

        # Verify files were created
        print("\nVerifying created files...")
        files_to_check = [
            (common_en, "Common EN textblock"),
            (common_de, "Common DE textblock"),
            (product_en, "Product EN textblock"),
            (product_de, "Product DE textblock"),
            (self.template_file_en, "EN template"),
            (self.template_file_de, "DE template"),
            (self.config_file, "Config file")
        ]

        for file_path, desc in files_to_check:
            if file_path.exists():
                print(f"✓ {desc} created at: {file_path}")
                print(f"  Size: {file_path.stat().st_size} bytes")
            else:
                print(f"✗ {desc} NOT created at: {file_path}")

        # Assert files exist
        for file_path, desc in files_to_check:
            self.assertTrue(file_path.exists(), f"{desc} was not created")

        print("\nTest files created:")
        print(f"Templates directory: {self.templates_dir}")
        print(f"Common textblocks: {self.textblocks_dir}/common")
        print(f"Product textblocks: {self.textblocks_dir}/products/{self.product_name}")
        print(f"Config file: {self.config_file}")

if __name__ == '__main__':
    unittest.main()
