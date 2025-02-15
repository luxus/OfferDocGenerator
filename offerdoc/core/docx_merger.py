from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.shared import Pt, RGBColor
from docxcompose.composer import Composer
from typing import Optional, List, Tuple, Dict
from pathlib import Path
from copy import deepcopy
import logging

logger = logging.getLogger(__name__)

class DocxMerger:
    def __init__(self, base_path: Path):
        self.base_doc = Document(base_path)
        # Track section numbers to maintain continuity
        self.section_map: Dict[str, int] = {}
        
        # Initialize section map based on existing sections in the base document
        self._initialize_section_map()
        
    def _initialize_section_map(self) -> None:
        """Scan base document and populate section map with existing headings."""
        for para in self.base_doc.paragraphs:
            text = para.text.strip()
            if self._is_heading(text):
                level = len(text.split('.')[:-1])
                key = '.'.join(text.split('.')[:-1])  # e.g., '1.2' from '1.2 Details'
                self.section_map[key] = int(text.split('.')[-1])

    def _is_heading(self, text: str) -> bool:
        """Check if a paragraph is a section heading."""
        return '.' in text and all(part.isdigit() for part in text.split('.'))

    def merge_content(
        self,
        source_path: Path,
        target_section: str,
        new_subsection: Optional[str] = None
    ) -> None:
        """
        Merge content from source DOCX into the base document at the specified section.
        
        Args:
            source_path: Path to source DOCX file
            target_section: Section heading where content should be inserted (e.g., '1.2 Details')
            new_subsection: Optional new subsection heading to create (e.g., '1.3 Additional Information')
        """
        try:
            source_doc = Document(source_path)
            
            # Find insertion point based on target section
            target_para = self._find_section(target_section)
            if not target_para:
                raise ValueError(f"Section '{target_section}' not found in base document")

            # Determine where to insert new content
            insert_position = None
            if hasattr(target_para, 'next_sibling'):
                insert_position = target_para.next_sibling
            if not insert_position:
                insert_position = target_para

            # Process each element from source document
            for element in source_doc.element.body:
                self._process_element(element, insert_position)

            # If creating a new subsection, add it before processing
            if new_subsection:
                self.create_new_section(new_subsection, after=target_section)

        except PackageNotFoundError as e:
            raise ValueError(f"Invalid DOCX file: {source_path}") from e
        except Exception as e:
            logger.error(f"Error merging files: {e}")
            raise

    def create_new_section(self, heading: str, after: Optional[str] = None) -> None:
        """
        Create a new section with the specified heading at the correct position.
        
        Args:
            heading: New section heading (e.g., '1.3 Additional Information')
            after: Section to insert new section after
        """
        # Determine the numbering for the new section
        sections = sorted(self.section_map.keys())
        if not sections:
            new_number = "1"
        else:
            last_section = max(sections)
            last_parts = list(map(int, last_section.split('.')))
            new_parts = [int(p) + 1 if i == len(last_parts)-1 else int(p) 
                        for i, p in enumerate(last_parts)]
            new_number = '.'.join(map(str, new_parts))
            
        # Create the new heading
        para = self.base_doc.add_paragraph()
        para.text = f"{new_number} {heading}"
        
        # Insert after specified section if provided
        if after:
            target_para = self._find_section(after)
            if not target_para:
                raise ValueError(f"Section '{after}' not found in base document")
            parent = target_para.parent
            parent.insert(parent.index(target_para) + 1, para._element)

    def _find_section(self, heading: str) -> Optional[BaseOxmlElement]:
        """Find the section element with the specified heading."""
        for para in self.base_doc.paragraphs:
            if para.text.strip() == heading:
                return para._element
        return None

    def _process_element(
        self,
        element: BaseOxmlElement,
        insert_position: Optional[BaseOxmlElement]
    ) -> None:
        """Process and insert an element from the source document."""
        if not insert_position:
            return

        try:
            parent = insert_position.parent
            if not parent:
                logger.warning("No parent found for insert position")
                return
                
            # Handle paragraph elements with their formatting
            if element.tag.endswith('p'):
                new_para = self.base_doc.add_paragraph()
                new_para._element.append(deepcopy(element))
                
                # Adjust font properties if needed
                self._apply_formatting(new_para)
                
                try:
                    index = parent.index(insert_position)
                    parent.insert(index, new_para._element)
                except (ValueError, IndexError):
                    # Fallback: append to end if insertion fails
                    parent.append(new_para._element)
        elif element.tag.endswith('tbl'):
            # Handle tables with their content
            new_table = self.base_doc.add_table(rows=0, cols=1)
            for row in element.findall('.//w:tr'):
                new_row = new_table.add_row()
                for cell in row.findall('.//w:tc'):
                    new_cell = new_row.cells[0]
                    self._copy_table_content(cell, new_cell)
        elif element.tag.endswith('sectPr'):
            # Handle section properties
            insert_position.append(deepcopy(element))

    def _apply_formatting(self, paragraph) -> None:
        """Apply formatting rules to maintain consistency."""
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            if 'bold' in run.text.lower():
                run.bold = True
            if 'italic' in run.text.lower():
                run.italic = True

    def _copy_table_content(self, source_cell: BaseOxmlElement, target_cell) -> None:
        """Copy content from a source cell to a target cell."""
        for element in source_cell.findall('.//w:p'):
            new_para = target_cell.add_paragraph()
            new_para._element.append(deepcopy(element))
            self._apply_formatting(new_para)

    def save(self, output_path: Path) -> None:
        """Save the merged document to the specified path."""
        try:
            self.base_doc.save(output_path)
            logger.info(f"Merged document saved to {output_path}")
        except Exception as e:
            logger.error(f"Error saving document: {e}")
            raise

    def _clean_up(self) -> None:
        """Clean up any temporary files or resources."""
        pass
