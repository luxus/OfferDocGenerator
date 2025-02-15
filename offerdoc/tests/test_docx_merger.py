import pytest
from pathlib import Path
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from ..core.docx_merger import DocxMerger
from ..core.exceptions import DocumentGenerationError

@pytest.fixture
def base_template(tmp_path):
    """Create a base template with some existing sections"""
    doc = Document()
    doc.add_heading('1. Introduction', level=1)
    doc.add_heading('1.1 Overview', level=2)
    doc.add_paragraph('Base template content')
    doc.add_heading('1.2 Details', level=2)
    doc.add_paragraph('More base content')
    
    base_path = tmp_path / "base.docx"
    doc.save(str(base_path))
    return base_path

@pytest.fixture
def source_doc1(tmp_path):
    """Create first source document with formatted content"""
    doc = Document()
    heading = doc.add_heading('2.3 Additional Details', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Test content from Source 1.')
    run.bold = True
    
    source_path = tmp_path / "source1.docx"
    doc.save(str(source_path))
    return source_path

@pytest.fixture
def source_doc2(tmp_path):
    """Create second source document with bullet points"""
    doc = Document()
    doc.add_heading('2.4 More Information', level=2)
    p = doc.add_paragraph('Content 2.')
    doc.add_paragraph('• Item 1', style='List Bullet')
    doc.add_paragraph('• Item 2', style='List Bullet')
    
    source_path = tmp_path / "source2.docx"
    doc.save(str(source_path))
    return source_path

def test_single_file_merge(tmp_path, base_template, source_doc1):
    """Test merging a single file with formatted content"""
    merger = DocxMerger(base_template)
    merger.merge_content(source_doc1, "1.2 Details")
    
    output_path = tmp_path / "merged_single.docx"
    merger.save(output_path)
    
    # Verify the merged document
    doc = Document(output_path)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert '2.3 Additional Details' in headings
    
    # Verify content and formatting
    content = [p.text for p in doc.paragraphs]
    assert 'Test content from Source 1.' in content

def test_multiple_files_merge(tmp_path, base_template, source_doc1, source_doc2):
    """Test merging multiple files with different content types"""
    merger = DocxMerger(base_template)
    merger.merge_content(source_doc1, "1.2 Details")
    merger.merge_content(source_doc2, "1.2 Details")
    
    output_path = tmp_path / "merged_multiple.docx"
    merger.save(output_path)
    
    # Verify the merged document
    doc = Document(output_path)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert '2.3 Additional Details' in headings
    assert '2.4 More Information' in headings
    
    # Verify bullet points
    content = [p.text for p in doc.paragraphs]
    assert '• Item 1' in content
    assert '• Item 2' in content

def test_error_handling(tmp_path, base_template):
    """Test error handling for missing and invalid files"""
    merger = DocxMerger(base_template)
    
    # Test missing file
    with pytest.raises(FileNotFoundError):
        merger.merge_content(tmp_path / "nonexistent.docx", "1.2 Details")
    
    # Test invalid target section
    with pytest.raises(ValueError):
        merger.merge_content(source_doc1, "Nonexistent Section")

def test_subsection_handling(tmp_path, base_template):
    """Test handling of nested subsections"""
    doc = Document()
    doc.add_heading('1.2.3 Sub-Detail', level=3)
    doc.add_paragraph('New sub-detail content')
    
    subsection_path = tmp_path / "subsection.docx"
    doc.save(str(subsection_path))
    
    merger = DocxMerger(base_template)
    merger.merge_content(subsection_path, "1.2 Details")
    
    output_path = tmp_path / "merged_subsection.docx"
    merger.save(output_path)
    
    # Verify subsection
    doc = Document(output_path)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]
    assert '1.2.3 Sub-Detail' in headings

def test_formatting_preservation(tmp_path, base_template):
    """Test preservation of text formatting"""
    doc = Document()
    heading = doc.add_heading('2.5 Formatted Section', level=2)
    
    p = doc.add_paragraph()
    run = p.add_run('Bold text. ')
    run.bold = True
    run = p.add_run('Italic text.')
    run.italic = True
    
    formatted_path = tmp_path / "formatted.docx"
    doc.save(str(formatted_path))
    
    merger = DocxMerger(base_template)
    merger.merge_content(formatted_path, "1.2 Details")
    
    output_path = tmp_path / "merged_formatted.docx"
    merger.save(output_path)
    
    # Verify formatting
    doc = Document(output_path)
    found_bold = False
    found_italic = False
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.bold:
                found_bold = True
            if run.italic:
                found_italic = True
    
    assert found_bold
    assert found_italic
