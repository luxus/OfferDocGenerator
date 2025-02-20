import sys
import os
import unittest
import shutil
import random
import zipfile
import zipfile
from pathlib import Path
from pathlib import Path
import yaml
import docx
from unittest import mock
from docxtpl import DocxTemplate, RichText
import offerdocgenerator

class TestOfferDocGenerator(unittest.TestCase):
    CLEANUP = False  # Set to False to keep generated files
    
    def test_directory_creation(self):
        """Verify test directory structure is created"""
        self.assertTrue(self.templates_dir.exists())
        self.assertTrue(self.textblocks_dir.exists())
        self.assertTrue((self.textblocks_dir / "common").exists())
        self.assertTrue((self.textblocks_dir / "products" / self.product_name).exists())

    CLEANUP = False  # Set to False to keep generated files
    TEST_DIR_NAME = "test_data"  # Fixed directory name
    
    def _validate_docx(self, path: Path) -> bool:
        """Validate DOCX file structure"""
        try:
            with zipfile.ZipFile(path) as z:
                if 'word/document.xml' not in z.namelist():
                    return False
                doc = docx.Document(path)
                return len(doc.paragraphs) > 0
        except (zipfile.BadZipFile, Exception):
            return False
            
    def setUp(self):
        """Set up test fixtures in shared directory"""
        # Base test directory
        self.script_dir = Path(__file__).parent
        self.test_root = self.script_dir / "test_output" 
        self.test_run_dir = self.test_root / self.TEST_DIR_NAME
        
        # Add test for Jinja2 loops and RichText
        self.loop_template = self.test_run_dir / "templates" / "loop_test.docx"
        
        # Clean existing test data if cleanup enabled
        if self.CLEANUP and self.test_run_dir.exists():
            shutil.rmtree(self.test_run_dir)
            
        # Create fresh directory structure
        self.test_run_dir.mkdir(parents=True, exist_ok=True)
        
        # Now create subdirectories
        self.config_file = self.test_run_dir / "test_config.yaml"
        self.templates_dir = self.test_run_dir / "templates"
        self.output_dir = self.test_run_dir / "output"
        self.textblocks_dir = self.test_run_dir / "textblocks"
        self.product_name = "Web Application Security Assessment"

        # Create required subdirectories
        self.templates_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)
        (self.textblocks_dir / "common").mkdir(parents=True, exist_ok=True)
        (self.textblocks_dir / "products" / self.product_name).mkdir(parents=True, exist_ok=True)

        # Create common textblocks
        self._create_textblock_file(
            self.textblocks_dir / "common" / "section_1_1_EN.docx",
            "Our standard security assessment provides a comprehensive evaluation of your web application's security posture."
        )
        self._create_textblock_file(
            self.textblocks_dir / "common" / "section_1_1_DE.docx",
            "Unsere Standard-Sicherheitsbewertung bietet eine umfassende Evaluation der Sicherheitslage Ihrer Webanwendung."
        )

        # Create product-specific textblocks for first product
        self._create_textblock_file(
            self.textblocks_dir / "products" / self.product_name / "section_1_1_1_EN.docx",
            """The Web Application Security Assessment includes:

- Vulnerability scanning
- Manual penetration testing
- Code review"""
        )
        self._create_textblock_file(
            self.textblocks_dir / "products" / self.product_name / "section_1_1_1_DE.docx",
            """Die Web Application Security Assessment beinhaltet:

- Schwachstellenscanning
- Manuelle Penetrationstests
- Code-Review"""
        )

        # Create API Security Review product and textblocks
        self.product_name2 = "API Security Review"
        (self.textblocks_dir / "products" / self.product_name2).mkdir(parents=True, exist_ok=True)
        self._create_textblock_file(
            self.textblocks_dir / "products" / self.product_name2 / "section_1_1_1_EN.docx",
            """The API Security Review includes:

- API endpoint security testing
- Authentication mechanism review
- Data validation assessment"""
        )
        self._create_textblock_file(
            self.textblocks_dir / "products" / self.product_name2 / "section_1_1_1_DE.docx",
            """Die API-Sicherheitsüberprüfung umfasst:

- API-Endpunkt-Sicherheitstests
- Überprüfung der Authentifizierungsmechanismen
- Bewertung der Datenvalidierung"""
        )

        # Create proper bundle templates with required variables
        self._create_bundle_templates()

        # Create base templates for EN and DE
        # English template
        self.template_file_en = self.templates_dir / "base_EN.docx"
        doc = docx.Document()
        doc.add_heading('Offer: {{ offer.number }}', 0)
        doc.add_paragraph('Date: {{ offer.date }}')
        doc.add_paragraph('Valid for: {{ offer.validity[LANGUAGE] }}')
        doc.add_heading('Customer Information', 1)
        doc.add_paragraph('{{ customer.name }}')
        doc.add_paragraph('{{ customer.address }}')
        doc.add_paragraph('{{ customer.city }}, {{ customer.zip }}')
        doc.add_paragraph('{{ customer.country }}')
        doc.add_heading('Product Description', 1)
        p = doc.add_paragraph()
        p.add_run('{{r section_1_1 }}')
        doc.add_heading('Detailed Scope', 2)
        p = doc.add_paragraph()
        p.add_run('{{r section_1_1_1 }}')
        doc.add_paragraph('Total Price: {{ CURRENCY }} 10,000')
        doc.add_heading('Sales Contact', 1)
        doc.add_paragraph('{{ sales.name }}')
        doc.add_paragraph('{{ sales.email }}')
        doc.add_paragraph('{{ sales.phone }}')
        doc.save(str(self.template_file_en))

        # German template
        self.template_file_de = self.templates_dir / "base_DE.docx"
        doc = docx.Document()
        doc.add_heading('Angebot: {{ offer.number }}', 0)
        doc.add_paragraph('Datum: {{ offer.date }}')
        doc.add_paragraph('Gültig für: {{ offer.validity[LANGUAGE] }}')
        doc.add_heading('Kundeninformationen', 1)
        doc.add_paragraph('{{ customer.name }}')
        doc.add_paragraph('{{ customer.address }}')
        doc.add_paragraph('{{ customer.city }}, {{ customer.zip }}')
        doc.add_paragraph('{{ customer.country }}')
        doc.add_heading('Produktbeschreibung', 1)
        p = doc.add_paragraph()
        p.add_run('{{r section_1_1 }}')
        doc.add_heading('Detaillierter Umfang', 2)
        p = doc.add_paragraph()
        p.add_run('{{r section_1_1_1 }}')
        doc.add_paragraph('Gesamtpreis: {{ CURRENCY }} 10.000')
        doc.add_heading('Vertriebskontakt', 1)
        doc.add_paragraph('{{ sales.name }}')
        doc.add_paragraph('{{ sales.email }}')
        doc.add_paragraph('{{ sales.phone }}')
        doc.save(str(self.template_file_de))

        # Create config file
        config = {
            "offer": {
                "number": "2025-001",
                "date": "2025-02-02",
                "validity": {
                    "EN": "30 days",
                    "DE": "30 Tage"
                }
            },
            "settings": {
                "products": str(self.textblocks_dir / "products"),
                "common": str(self.textblocks_dir / "common"),
                "output": str(self.output_dir),
                "templates": str(self.templates_dir),
                "format": "docx",
                "prefix": "TestOffer_"
            },
            "customer": {
                "name": "Example Corp",
                "address": "123 Example Street",
                "city": "Example City",
                "zip": "12345",
                "country": "Example Country"
            },
            "sales": {
                "name": "John Doe",
                "email": "john.doe@example.com",
                "phone": "+1 234 567 890"
            },
            "bundles": {
                "web_security_pack": {
                    "name": "Web Security Package",
                    "products": ["Web Application Security Assessment", "API Security Review"],
                    "discount": {"percentage": 15.0},
                    "template": "bundle_base_{language}.docx",
                    "variables": {}
                }
            }
        }
        with open(self.config_file, 'w') as f:
            yaml.dump(config, f)

    def _create_textblock_file(self, file_path: Path, content: str):
        """Helper method to create a docx file with given content."""
        doc = docx.Document()
        # Add each paragraph with proper styling
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                p = doc.add_paragraph()
                # If it's a bullet point, use a list style
                if paragraph.strip().startswith('-'):
                    p.style = 'List Bullet'
                    p.text = paragraph.strip()[2:]  # Remove the '- ' prefix
                else:
                    p.text = paragraph.strip()
        doc.save(str(file_path))

    def tearDown(self):
        """Conditional cleanup of test output directory"""
        if self.CLEANUP:
            shutil.rmtree(self.test_run_dir)
        else:
            print(f"\nTest files preserved in: {self.test_run_dir}")

    def test_get_product_names(self):
        """Test that product names are correctly detected from the directory structure."""
        config = offerdocgenerator.load_config(self.config_file)
        products = offerdocgenerator.get_product_names(config)
        self.assertEqual(len(products), 2)
        self.assertIn(self.product_name, products)
        self.assertIn(self.product_name2, products)

    def test_load_config(self):
        """Test that the YAML configuration is loaded correctly."""
        config = offerdocgenerator.load_config(self.config_file)
        self.assertIsInstance(config, offerdocgenerator.Config)
        self.assertEqual(config.settings.products, self.textblocks_dir / "products")
        self.assertEqual(config.settings.common, self.textblocks_dir / "common")

    def test_load_textblocks(self):
        """Test dynamic loading of textblocks from product directory"""
        config = offerdocgenerator.load_config(self.config_file)
        
        # Create a template instance
        template = DocxTemplate(str(self.template_file_en))
        
        # Test German textblocks
        section_1_1_de, _ = offerdocgenerator.load_textblock("section_1_1", config, self.product_name, "DE", template)
        section_1_1_1_de, _ = offerdocgenerator.load_textblock("section_1_1_1", config, self.product_name, "DE", template)
        
        self.assertIsNotNone(section_1_1_de)
        self.assertIsNotNone(section_1_1_1_de)
        self.assertIn("Sicherheitsbewertung", str(section_1_1_de))
        self.assertIn("Schwachstellenscanning", str(section_1_1_1_de))
        
        # Test English textblocks 
        section_1_1_en, _ = offerdocgenerator.load_textblock("section_1_1", config, self.product_name, "EN", template)
        section_1_1_1_en, _ = offerdocgenerator.load_textblock("section_1_1_1", config, self.product_name, "EN", template)
        
        self.assertIsNotNone(section_1_1_en)
        self.assertIsNotNone(section_1_1_1_en)
        self.assertIn("comprehensive evaluation", str(section_1_1_en))
        self.assertIn("Vulnerability scanning", str(section_1_1_1_en))

    def test_template_variable_detection(self):
        """Test that template variables are properly detected"""
        # Create a test template in the temporary directory
        test_template = self.templates_dir / "test.docx"
        doc = docx.Document()
        doc.add_paragraph('{{ sales_name }}')
        doc.add_paragraph('{{ sales_email }}')
        doc.add_paragraph('{{ sales_phone }}')
        doc.save(str(test_template))
        
        # Load template and detect variables
        doc = DocxTemplate(str(test_template))
        detected_vars = doc.get_undeclared_template_variables()
        
        # Verify expected variables are detected
        required_vars = {'sales_name', 'sales_email', 'sales_phone'}
        self.assertTrue(required_vars.issubset(detected_vars),
                       f"Missing sales fields in template variables: {required_vars - detected_vars}")

    def test_error_handling(self):
        """Test error handling for missing files"""
        with self.assertRaises(SystemExit):
            offerdocgenerator.load_config(Path("/non/existent/config.yaml"))

    def test_jinja_loops_and_richtext(self):
        """Test Jinja2 loops with config data and RichText formatting"""
        # Create PROPERLY STRUCTURED template with loop in single paragraph
        doc = docx.Document()
        
        # Create a single paragraph for the entire loop structure
        p = doc.add_paragraph()
        runner = p.add_run()
        runner.add_text("Contacts:\n")
        
        # Add loop control in the same run with line breaks
        runner.add_text("{% for contact in sales.contacts %}\n")
        runner.add_text("{{ contact.name }} - {{ contact.email }}\n")
        runner.add_text("{% endfor %}")
        
        # Add RichText field in separate paragraph
        doc.add_paragraph().add_run("Customer: {{r customer.name }}")
        
        doc.save(str(self.loop_template))
        
        # Update config with ALL required fields
        config_data = {
            'customer': {
                'name': 'Example Corp',
                'address': '123 Example Street',
                'city': 'Example City',
                'zip': '12345',
                'country': 'Example Country'
            },
            'offer': {
                'number': '2023-TEST',
                'date': '2023-01-01',
                'validity': {
                    'EN': '30 days',
                    'DE': '30 Tage'
                }
            },
            'settings': {
                'products': 'products',
                'common': 'common',
                'output': 'output',
                'templates': 'templates',
                'format': 'docx',
                'filename_pattern': 'Offer_{product}.docx',
                'template_pattern': 'base_{language}.docx'
            },
            'sales': {
                'name': 'Test Sales',
                'email': 'sales@example.com',
                'phone': '+1234567890',
                'contacts': [
                    {'name': 'Alice', 'email': 'alice@example.com'},
                    {'name': 'Bob', 'email': 'bob@example.com'}
                ]
            },
            'languages': ["EN", "DE"],
            'currencies': ["USD", "EUR"],
            'textblock_patterns': [
                "{var_name}_{language}.docx",
                "{var_name}.docx"
            ]
        }
    
        with open(self.config_file, 'w') as f:
            yaml.dump(config_data, f)
        
        # Load config properly using load_config
        config = offerdocgenerator.load_config(self.config_file)  # Get Config instance
    
        # Render document with the loaded config
        template = DocxTemplate(str(self.loop_template))
        context = offerdocgenerator.build_context(config, "EN", self.product_name, "CHF")
        output_path = self.output_dir / "loop_test.docx"
        offerdocgenerator.render_offer(template, config, context, output_path)
        
        # Verify output
        doc = docx.Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        
        # Test loop results
        self.assertIn("Alice - alice@example.com", full_text)
        self.assertIn("Bob - bob@example.com", full_text)
        
        # Test RichText formatting
        with zipfile.ZipFile(output_path) as z:
            content = z.read("word/document.xml").decode()
            self.assertIn("Example Corp", content)


    @unittest.skip("Temporarily disabled - needs investigation of DOTX template handling")
    def test_dotx_generation(self):
        """Verify DOTX template creation with correct content type
        TODO: Investigate proper way to test DOTX template generation and usage
        Current issue: DocxTemplate cannot properly handle DOTX files
        Possible solutions:
        1. Use different library for DOTX testing
        2. Modify test to verify content type without opening template
        3. Convert DOTX to DOCX before testing content
        """
        config = offerdocgenerator.load_config(self.config_file)
        
        # Override output format for this test
        config.output["format"] = "dotx"
        product = self.product_name
        lang = "EN"
        currency = "EUR"
        
        # Generate DOTX file
        context = offerdocgenerator.build_context(config, lang, product, currency)
        textblocks = offerdocgenerator.load_textblocks(config, config.offer["sections"], product, lang)
        context.update(textblocks)
        
        output_file = self.output_dir / product / f"TEST_DOTX_{product}_{lang}_{currency}.dotx"
        output_file.parent.mkdir(parents=True, exist_ok=True)
        offerdocgenerator.render_offer(self.template_file_en, context, output_file)
        
        # Verify file properties
        self.assertTrue(output_file.exists())
        
        # Check ZIP package content type
        with zipfile.ZipFile(output_file) as z:
            content_type = z.read('[Content_Types].xml').decode()
            self.assertIn('application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml', content_type)
        
        # Verify template can be used to create new documents
        test_output = output_file.with_name("test_from_template.docx")
        try:
            # Create new document FROM TEMPLATE using docxtpl
            template = DocxTemplate(str(output_file))
            template.render({"test_content": "Test added via template"})
            template.save(str(test_output))
            self.assertTrue(test_output.exists())
            
            # Verify the new document is valid
            test_doc = docx.Document(str(test_output))
            self.assertIn("Test added via template", test_doc.paragraphs[0].text)
        finally:
            if test_output.exists():
                test_output.unlink()

    def test_special_characters_in_templates(self):
        """Test templates with special characters"""
        # Create test template with special chars
        special_template = self.templates_dir / "special_chars.docx"
        doc = docx.Document()
        doc.add_paragraph('Test & Company © 2024')
        doc.add_paragraph('{{ sales.email }}')
        doc.save(str(special_template))
        
        # Render and verify
        config = offerdocgenerator.load_config(self.config_file)
        context = offerdocgenerator.build_context(config, "EN", self.product_name, "CHF")
        output_path = self.output_dir / "special_chars_test.docx"
        
        template = DocxTemplate(str(special_template))
        offerdocgenerator.render_offer(template, config, context, output_path)
        
        # Verify output
        doc = docx.Document(str(output_path))
        self.assertIn("Test & Company © 2024", doc.paragraphs[0].text)
        self.assertIn("john.doe@example.com", doc.paragraphs[1].text)

    def test_bundle_template_processing(self):
        """Test end-to-end bundle document generation with actual template"""
        # Load config with bundle settings
        config = offerdocgenerator.load_config(self.config_file)
        
        # Generate bundle offers
        offerdocgenerator.generate_bundle_offer(config, "web_security_pack")
        
        # Verify output file
        output_path = (
            config.output_path / "bundles" / "web_security_pack" / 
            f"Offer_Web Security Package_EN_CHF_{config.offer.date}.docx"
        )
        self.assertTrue(output_path.exists())
        
        # Verify template content
        doc = docx.Document(str(output_path))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        
        # Check required bundle elements
        self.assertIn("Bundle Package: Web Security Package", full_text)
        self.assertIn("Bundle Discount: 15%", full_text)
        self.assertIn("Web Application Security Assessment", full_text)
        self.assertIn("API Security Review", full_text)

    def test_generated_file_validity(self):
        """Verify all generated files are valid Word documents"""
        config = offerdocgenerator.load_config(self.config_file)
        
        # Generate test documents
        for lang in ["EN", "DE"]:
            for currency in ["CHF", "EUR"]:
                context = offerdocgenerator.build_context(config, lang, self.product_name, currency)
                template = DocxTemplate(str(self.template_file_en if lang == "EN" else self.template_file_de))
                output_path = self.output_dir / f"test_doc_{lang}_{currency}.docx"
                offerdocgenerator.render_offer(template, config, context, output_path)
                
                # Verify file validity
                self.assertTrue(self._validate_docx(output_path), 
                              f"Invalid DOCX file: {output_path.name}")
                
                # Check content
                doc = docx.Document(output_path)
                full_text = "\n".join(p.text for p in doc.paragraphs)
                self.assertIn(currency, full_text)
                self.assertIn(config.offer.number, full_text)
                
    def test_bundle_template_variables(self):
        """Verify required variables exist in bundle template"""
        template_path = self.templates_dir / "bundle_base_EN.docx"
        doc = DocxTemplate(str(template_path))
        variables = doc.get_undeclared_template_variables()
        
        required_vars = {
            'bundle',
            'products',
            'discount'
        }
        missing = required_vars - set(variables)
        self.assertEqual(missing, set(), 
                        f"Missing variables: {missing}. Found: {variables}")

    def test_bundle_template_security(self):
        """Test bundle template security constraints"""
        template_path = self.templates_dir / "bundle_base_EN.docx"
        
        # Size check
        max_size = 10 * 1024 * 1024  # 10MB limit
        self.assertLess(template_path.stat().st_size, max_size, "Template too large")
        
        # Ownership check
        self.assertEqual(template_path.owner(), Path(__file__).owner(), "Owner mismatch")
        
        # Content validation
        doc = docx.Document(str(template_path))
        content = "\n".join(p.text for p in doc.paragraphs)
        forbidden_patterns = [
            r'\{\{.*\.(save|delete|write).*\}\}',
            r'\{\{.*__.*\}\}',
            r'\{\{.*config\.security.*\}\}'
        ]
        
        for pattern in forbidden_patterns:
            self.assertNotRegex(content, pattern, "Forbidden pattern found")

    def test_richtext_format_preservation(self):
        """Verify bold/italic/underline formatting in textblocks"""
        # Create formatted test file with correct section prefix
        test_file = self.textblocks_dir / "common" / "section_formatted_EN.docx"
        doc = docx.Document()
        p = doc.add_paragraph()
        p.add_run('Bold text').bold = True
        p.add_run(' Italic').italic = True
        p.add_run(' Underlined').underline = True
        doc.save(str(test_file))

        # Load and verify textblocks
        config = offerdocgenerator.load_config(self.config_file)
        
        # Create template instance
        template = DocxTemplate(str(self.template_file_en))
        
        rt, _ = offerdocgenerator.load_textblock("section_formatted", config, self.product_name, "EN", template)
        self.assertIsNotNone(rt)  # Verify section exists
        
        # Render document to check formatting
        output_path = self.output_dir / "richtext_test.docx"
        
        # Create template with proper section variable
        doc = docx.Document(str(self.template_file_en))
        p = doc.add_paragraph()
        p.add_run('{{r section_formatted }}')
        doc.save(str(self.template_file_en))
        
        # Create template and render with proper context
        template = DocxTemplate(str(self.template_file_en))
        context = offerdocgenerator.build_context(config, "EN", self.product_name, "CHF")
        context["section_formatted"] = rt
        template.render(context)
        template.save(str(output_path))
        
        # Check formatting in rendered document
        with zipfile.ZipFile(output_path) as z:
            document_xml = z.read("word/document.xml").decode()
            self.assertIn("<w:b/>", document_xml)  # Bold tag
            self.assertIn("<w:i/>", document_xml)  # Italic tag
            self.assertIn('<w:u w:val="single"/>', document_xml)  # Underline tag

    def test_invalid_config_sections(self):
        """Test handling of config with missing required sections"""
        invalid_config = {
            "offer": {"number": "123"},  # Missing settings, customer, sales
        }
        
        with open(self.config_file, 'w') as f:
            yaml.dump(invalid_config, f)
        
        with self.assertRaises(ValueError) as cm:
            offerdocgenerator.load_config(self.config_file)
        error_msg = str(cm.exception)
        self.assertIn("field required", error_msg.lower())
        for section in ['settings', 'customer', 'sales']:
            self.assertIn(section, error_msg.lower())

    def test_invalid_config_fields(self):
        """Test missing required fields within existing sections"""
        invalid_config = {
            "offer": {
                "number": "123",
                # Missing date and validity
            },
            "settings": {
                "products": "./products",
                # Missing common, output, and template_prefix
            },
            "customer": {
                "name": "Test Corp",
                "address": "Test St",
                "city": "Test City",
                "zip": "12345",
                "country": "Test Country"
            },
            "sales": {
                "name": "Test Sales",
                "email": "test@example.com",
                "phone": "123456"
            }
        }
        
        with open(self.config_file, 'w') as f:
            yaml.dump(invalid_config, f)
        
        with self.assertRaises(ValueError) as cm:
            offerdocgenerator.load_config(self.config_file)
        error_msg = str(cm.exception)
        # Check for Pydantic validation errors
        self.assertIn("offer.date\n  Field required", error_msg)
        self.assertIn("offer.validity\n  Field required", error_msg)
        self.assertIn("settings.common\n  Field required", error_msg)
        self.assertIn("settings.output\n  Field required", error_msg)
        self.assertIn("settings.templates\n  Field required", error_msg)

    def test_custom_settings_with_defaults(self):
        """Verify custom settings override defaults and missing settings use defaults."""
        custom_config = {
            "offer": {
                "number": "2025-002",
                "date": "2025-03-03",
                "validity": {
                    "EN": "45 days",
                    "DE": "45 Tage"
                }
            },
            "settings": {
                "products": "./custom_products",
                "common": "./custom_common",
                "output": "./custom_output",
                "templates": "custom_template",
                # format and prefix will use defaults
            },
            "customer": {
                "name": "Test Corp",
                "address": "456 Test Ave",
                "city": "Testville",
                "zip": "67890",
                "country": "Testland"
            },
            "sales": {
                "name": "Jane Smith",
                "email": "jane.smith@example.com",
                "phone": "+44 987 654 321"
            }
        }
        
        custom_config_path = self.test_run_dir / "custom_config.yaml"
        with open(custom_config_path, 'w') as f:
            yaml.dump(custom_config, f)
        
        config = offerdocgenerator.load_config(custom_config_path)
        
        # Verify custom settings with resolved paths
        self.assertEqual(config.settings.products, Path(custom_config_path.parent / "custom_products").resolve())
        self.assertEqual(config.settings.common, Path(custom_config_path.parent / "custom_common").resolve())
        self.assertEqual(config.settings.output, Path(custom_config_path.parent / "custom_output").resolve())
        self.assertEqual(config.settings.templates, Path(custom_config_path.parent / "custom_template").resolve())
        
        # Verify defaults
        self.assertEqual(config.settings.format, "docx")
        self.assertEqual(getattr(config.settings, "prefix", "Offer_"), "Offer_")
        
        # Verify other sections loaded correctly
        self.assertEqual(config.offer.number, "2025-002")
        self.assertEqual(config.customer.city, "Testville")
        self.assertEqual(config.sales.name, "Jane Smith")
        
        # Verify template path construction
        template_path_en = Path("custom_template") / "base_EN.docx"
        template_path_de = Path("custom_template") / "base_DE.docx"
        expected_path_en = (self.test_run_dir / "custom_template" / "base_EN.docx").resolve()
        expected_path_de = (self.test_run_dir / "custom_template" / "base_DE.docx").resolve()
        self.assertEqual(config.settings.templates, expected_path_en.parent)
        self.assertEqual(config.settings.templates, expected_path_de.parent)

    def test_render_offer(self):
        """Test rendering for all language/currency combinations in both DOCX and DOTX formats."""
        # Load fresh config for each format test
        for output_format in ["docx", "dotx"]:
            config = offerdocgenerator.load_config(self.config_file)
            from copy import deepcopy
            new_settings = deepcopy(config.settings)
            new_settings.format = output_format
            config.settings = new_settings
            
            products = offerdocgenerator.get_product_names(config)
            prefix = "TestOffer_"  # Match the prefix set in setUp()

            # Add validity text to templates for nested config testing
            for template in [self.template_file_en, self.template_file_de]:
                doc = docx.Document(str(template))
                doc.add_paragraph('Validity: {{ offer.validity[LANGUAGE] }}')
                doc.save(str(template))

            # Verify files per product
            for product in products:
                for lang in ["EN", "DE"]:
                    for currency in ["CHF", "EUR"]:
                        with self.subTest(product=product, language=lang, currency=currency, format=output_format):
                            # Build context with currency using build_context
                            context = offerdocgenerator.build_context(config, lang, product, currency)
            
                            # Select and load template
                            template_path = self.template_file_en if lang == "EN" else self.template_file_de
                            template = DocxTemplate(str(template_path))
            
                            # Get template variables and resolve them
                            template_vars = template.get_undeclared_template_variables()
                            # Filter variables not already in context
                            vars_to_resolve = {var for var in template_vars if var not in context}
                            resolved = offerdocgenerator.resolve_template_variables(vars_to_resolve, config, product, lang, template)
                            context.update(resolved)
            
                            # Generate output path using config properties
                            output_file = config.settings.output_path / product / f"{prefix}{product}_{lang}_{currency}.{output_format}"
                            
                            # Create product subdirectory
                            output_file.parent.mkdir(parents=True, exist_ok=True)
                            
                            # Render and verify
                            offerdocgenerator.render_offer(template, config, context, output_file)
                            self.assertTrue(output_file.exists())
                            
                            # Only validate DOCX content - skip for DOTX
                            if output_format == "docx":
                                doc = docx.Document(str(output_file))
                                full_text = "\n".join(para.text for para in doc.paragraphs)
                                self.assertIn(currency, full_text)
                                self.assertIn(config.offer.number, full_text)
                                self.assertIn(config.customer.name, full_text)
                                self.assertIn(config.customer.address, full_text)
                                self.assertIn(config.sales.email, full_text)
                                self.assertIn(config.sales.phone, full_text)

            # Verify file count for this format
            generated_files = list(self.output_dir.glob(f"**/{prefix}*.{output_format}"))
            self.assertEqual(len(generated_files), 8,
                           f"Expected 8 files for {output_format}, found {len(generated_files)}")

    def _create_bundle_templates(self):
        """Generate test bundle templates programmatically"""
        for lang in ['EN', 'DE']:
            doc = docx.Document()
            doc.add_paragraph('{{ bundle.name }}')
            doc.add_paragraph('Bundle Package: {{ bundle.name }}')
            doc.add_paragraph('Bundle Discount: {{ discount }}%')
            doc.add_paragraph('Products: {% for product in products %}{{ product }}{% endfor %}')
            template_path = self.templates_dir / f"bundle_base_{lang}.docx"
            doc.save(template_path)

if __name__ == '__main__':
    unittest.main()
