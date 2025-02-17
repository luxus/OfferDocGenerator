import os
from pathlib import Path
import pytest
import yaml
from docx import Document
from odg.utils.file_handler import FileHandler
from odg.document import DocumentSection, SectionType, DocumentMode
from odg.main import ConfigGenerator
from config.settings import Config, FolderConfig

@pytest.fixture
def config_generator(tmp_path):
    """Fixture providing a ConfigGenerator instance with a temporary directory"""
    test_dir = tmp_path / "test_project"
    test_dir.mkdir(parents=True, exist_ok=True)
    
    folders = FolderConfig(
        templates=test_dir / "templates",
        common=test_dir / "common",
        products=test_dir / "products",
        output=test_dir / "output",
        generated=test_dir / "generated"
    )
    
    config = Config(
        base_path=test_dir,
        output_dir=test_dir,
        folders=folders
    )
    
    return ConfigGenerator(output_dir=test_dir)

def test_generate_base_template_with_jinja_variables(config_generator, tmp_path):
    """Test base template generation with numbered lists"""
    # First create the base template
    template_name = "base_en.docx"
    template_path = config_generator.create_docx_template(template_name)
    assert template_path.exists()
    
    # Then create the sample document
    product_name = "SampleProduct"
    language = "en"
    currency = "USD"
    output_path = config_generator.create_sample_docx(product_name, language, currency)
    assert output_path.exists()
    
    # Verify Jinja2 variables were replaced
    with open(output_path, 'rb') as doc_file:
        doc = Document(doc_file)
        doc_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        assert '{{' not in doc_text, "Unrendered Jinja2 variables found in document"
    
    # Check for nested numbered lists
    file_handler = FileHandler(Config())
    assert file_handler.has_numbered_lists(output_path), "Document should contain nested numbered lists"
    
def test_generate_product_templates_with_numbered_lists(config_generator, tmp_path):
    """Test product template generation with numbered lists"""
    # First create the base template
    template_name = "base_en.docx"
    template_path = config_generator.create_docx_template(template_name)
    assert template_path.exists()
    
    product_name = "Web Application Security Assessment"
    language = "en"
    currency = "EUR"
    output_path = config_generator.create_sample_docx(product_name, language, currency)
    assert output_path.exists()
    
    # Check for nested numbered lists
    document = Document(str(output_path))
    has_numbered_list = any(p.style.name == 'List Number' for p in document.paragraphs)
    has_nested_list = any(p.style.name == 'List Number 2' for p in document.paragraphs)
    assert has_numbered_list and has_nested_list
        
def test_validate_base_template(config_generator, tmp_path):
    """Verify base template structure and content"""
    # First create the base template
    template_name = "base_en.docx"
    template_path = config_generator.create_docx_template(template_name)
    assert template_path.exists()
    
    product_name = "SampleProduct"
    language = "en"
    currency = "USD"
    output_path = config_generator.create_sample_docx(product_name, language, currency)
    assert output_path.exists()
    
    file_handler = FileHandler(Config())
    assert file_handler.validate_merged_doc(output_path), f"Invalid document structure: {output_path}"

def test_merge_documents(config_generator, tmp_path):
    """Test merging multiple DOCX files with continuous numbering."""
    product_name = "Web Application Security Assessment"
    language = "en"
    currency = "EUR"
    
    # Create individual templates
    template1 = config_generator.create_docx_template("section1.docx")
    template2 = config_generator.create_docx_template("section2.docx")
    
    # Merge the documents
    merged_path = tmp_path / "merged.docx"
    config_generator.merge_docx_files([template1, template2], merged_path)
    
    file_handler = FileHandler(Config())
    assert file_handler.validate_merged_doc(merged_path), f"Merged document {merged_path} has invalid numbering"

def test_validate_product_template(config_generator, tmp_path):
    """Verify product template structure and content"""
    # First create the base template
    template_name = "base_en.docx"
    template_path = config_generator.create_docx_template(template_name)
    assert template_path.exists()
    
    product_name = "Web Application Security Assessment"
    language = "en"
    currency = "EUR"
    output_path = config_generator.create_sample_docx(product_name, language, currency)
    assert output_path.exists()
    
    file_handler = FileHandler(Config())
    assert file_handler.validate_doc_structure(output_path), f"Invalid document structure: {output_path}"
    assert file_handler.has_required_sections(output_path), f"Missing required sections: {output_path}"
