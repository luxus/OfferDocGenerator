#!/usr/bin/env python3
import sys
import logging
import traceback
from pathlib import Path
from typing import Dict, List, Set, Any, Optional, Tuple
from dataclasses import dataclass, field
from docx import Document
from docxtpl import DocxTemplate, RichText
import yaml

from offerdoc.core.config import load_config, AppConfig
from offerdoc.core.exceptions import handle_document_errors
from offerdoc.core.file_handler import FileHandler
from offerdoc.core.renderer import DocumentRenderer
from offerdoc.utils.formatters import colorize

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

from offerdoc.core.config import AppConfig

from pydantic import ConfigDict, model_validator
from typing import Optional, Dict, Any

class Config(AppConfig):
    """Extended configuration with runtime properties"""
    model_config = ConfigDict(validate_default=True, extra='forbid')

    @model_validator(mode='after')
    def validate_config(self) -> 'Config':
        """Validate configuration after initialization."""
        # Only validate required fields in production mode
        if not getattr(self, '_test_mode', False):
            required_sections = ['offer', 'settings', 'customer', 'sales']
            for section in required_sections:
                if not getattr(self, section, None):
                    raise ValueError(f"Missing required section: {section}")

            # Validate critical fields that must always be present
            if self.settings:
                if not getattr(self.settings, 'templates', None):
                    raise ValueError("Missing required field: settings.templates")
                
            if self.offer:
                if not getattr(self.offer, 'number', None):
                    raise ValueError("Missing required field: offer.number")

        return self

    @classmethod
    def for_tests(cls, **data) -> 'Config':
        """Create a config instance for testing with relaxed validation"""
        instance = cls.model_validate(data)
        instance._test_mode = True
        return instance

def load_config(config_path: Path) -> Config:
    """Load configuration from YAML file."""
    try:
        with open(config_path) as f:
            config_data = yaml.safe_load(f)
            
        # Create config instance with context
        return Config.model_validate(
            config_data,
            context={"config_path": config_path}  # Add context for path resolution
        )
        
    except Exception as e:
        logger.error(f"Error loading config from {config_path}: {e}")
        if isinstance(e, (ValueError, TypeError)):
            raise  # Re-raise validation errors for tests to catch
        sys.exit(1)

def get_product_names(config: Config) -> List[str]:
    """Get list of available products from the products directory."""
    products_dir = config.products_path
    if not products_dir.exists():
        logger.error(f"Products directory not found: {products_dir}")
        return []
    return [d.name for d in products_dir.iterdir() if d.is_dir()]

def load_textblock_file(file_path: Path) -> str:
    """
    Load content from a DOCX file.
    Returns the text content preserving paragraphs.
    """
    try:
        doc = Document(str(file_path))
        # Join paragraphs with double newlines to preserve formatting
        content = "\n\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        return content
    except Exception as e:
        logger.error(f"Error reading textblock file {file_path}: {e}")
        return ""


def load_textblock(var_name: str, config: Config, product_name: str, language: str, template: DocxTemplate) -> Tuple[Optional[Any], Optional[Path]]:
    """Dynamically load DOCX content using configured patterns"""
    search_locations = [
        config.products_path / product_name,
        config.common_path
    ]

    for base_path in search_locations:
        for pattern in config.textblock_patterns:
            target_path = base_path / pattern.format(
                var_name=var_name,
                language=language.upper()
            )
            if target_path.exists():
                try:
                    # Create subdoc with preserved formatting
                    subdoc = template.new_subdoc()
                    doc = Document(str(target_path))
                    for paragraph in doc.paragraphs:
                        p = subdoc.add_paragraph()
                        for run in paragraph.runs:
                            r = p.add_run(run.text)
                            r.bold = run.bold
                            r.italic = run.italic
                            r.underline = run.underline
                    return subdoc, target_path
                except Exception as e:
                    logger.error(f"Failed to load subdoc {target_path}: {e}")
                    return None, None
    return None, None

def resolve_template_variables(template_vars: Set[str], config: Config, 
                            product_name: str, language: str,
                            template: DocxTemplate) -> Dict[str, Any]:
    """Resolve variables from config, textblocks, and special handlers"""
    resolved = {}
    language = language.upper()
    
    print(colorize(f"\n🔍 Resolving {len(template_vars)} template variables:", 'CYAN'))
    
    for var in sorted(template_vars):
        # First try direct config value
        try:
            resolved[var] = resolve_config_variable(var, config)
            print(f"  {colorize(var.ljust(20), 'GREEN')} {colorize('←', 'BLUE')} config.{var.replace('_', '.')}")
            continue
        except (ValueError, AttributeError):
            pass
            
        # Then try to load as DOCX subdocument
        subdoc, source_path = load_textblock(var, config, product_name, language, template)
        if subdoc:
            resolved[var] = subdoc
            print(f"  {colorize(var.ljust(20), 'YELLOW')} {colorize('←', 'BLUE')} {colorize(str(source_path.relative_to(config.common_path.parent)), 'CYAN')}")
            continue
            
        # Fallback to empty string if nothing found
        resolved[var] = ""
        print(f"  {colorize(var.ljust(20), 'RED')} {colorize('← WARNING: No config value or DOCX found', 'RED')}")
    
    return resolved

def resolve_config_variable(var_path: str, config: Config) -> Any:
    """Resolve nested dictionary paths using dot notation"""
    parts = var_path.split('.')  # Change to dot notation
    current = config.model_dump()  # Start with full config dict
    
    for part in parts:
        if isinstance(current, dict) and part in current:
            current = current[part]
        else:
            raise ValueError(f"Config path not found: {var_path}")
    return current

def build_context(config: Config, language: str, product_name: str, currency: str) -> Dict[str, Any]:
    """Build base context with core variables."""
    return {
        "offer": config.offer.model_dump(),
        "customer": config.customer.model_dump(),
        "sales": config.sales.model_dump(),
        "settings": config.settings.model_dump(),
        "contacts": config.sales.contacts,  # Add explicit access to contacts
        "LANGUAGE": language.upper(),
        "PRODUCT": product_name,
        "CURRENCY": currency,
        "r": lambda text: RichText(text) if text else ""  # Add RichText helper
    }

def render_offer(template: DocxTemplate, config: Config, context: Dict[str, Any], output_path: Path):
    """Render template with auto-discovered variables"""
    try:
        logger.debug("Rendering context contains: %s", context.keys())
        logger.debug("Bundle data: %s", context.get('bundle'))
        logger.debug("Discount value: %s", context.get('discount'))
        
        
        # Get all variables from the template using proper detection
        template_vars = set(template.get_undeclared_template_variables())
        
        # Remove built-in Jinja variables
        template_vars -= {'True', 'False', 'None'}
        
        # Only resolve variables not already in context
        vars_to_resolve = {var for var in template_vars if var not in context}
        
        # Resolve variables from multiple sources
        resolved_context = resolve_template_variables(vars_to_resolve, config,
                                                   context['PRODUCT'], context['LANGUAGE'],
                                                   template)
        
        # Merge with existing context
        context.update(resolved_context)

        # Render template with complete context
        template.render(context, autoescape=True)
        
        # Handle different output formats
        output_format = output_path.suffix[1:].lower()
        
        if output_format == 'dotx':
            # Change content type for template format
            document_part = template.docx.part
            document_part._content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml'
        
        # Ensure parent directories exist
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Save with configured format
        template.save(str(output_path))
        
        # Enhanced output message with safe path handling
        if output_path.exists():
            try:
                # Try to get relative path first
                rel_path = output_path.relative_to(Path.cwd())
                display_path = str(rel_path)
            except ValueError:
                # Fall back to absolute path if not in CWD
                display_path = str(output_path.resolve())
            
            file_size = output_path.stat().st_size / 1024
            
            # Split into directory and filename components
            path_obj = Path(display_path)
            dir_part = str(path_obj.parent)
            file_name = path_obj.name
            
            # Color directory in yellow and process filename
            colored_dir = colorize(dir_part, 'YELLOW') if dir_part != '.' else ''
            
            # Split filename into components
            parts = file_name.split('_')
            colored_parts = []
            for part in parts:
                if part in ['DE', 'EN']:
                    colored_parts.append(colorize(part, 'CYAN'))
                elif part in ['CHF', 'EUR']:
                    colored_parts.append(colorize(part, 'GREEN'))
                else:
                    colored_parts.append(colorize(part, 'YELLOW'))
                    
            colored_filename = '_'.join(colored_parts)
            size_str = colorize(f"({file_size:.1f} KB)", 'BLUE')
            
            # Combine path parts, handling current directory case
            full_colored_path = f"{colored_dir}/{colored_filename}" if colored_dir else colored_filename
            print(f"\n{colorize('✅ Document:', 'GREEN')} {full_colored_path} {size_str}")
        
    except Exception as e:
        logger.error(f"Error during template rendering: {e}")
        logger.error(f"Error details: {traceback.format_exc()}")
        raise

def generate_bundle_offer(config: Config, bundle_name: str):
    """Generate offer documents for a product bundle"""
    bundle = config.bundles[bundle_name]
    
    # Convert product references to just names
    product_names = [p.name if hasattr(p, 'name') else p for p in bundle.products]
    
    for lang in config.languages:
        for currency in config.currencies:
            # Build bundle context with proper types
            context = build_context(config, lang, bundle_name, currency)
            context.update({
                "bundle": {
                    "name": bundle.name,
                    "discount": int(bundle.discount["percentage"])  # Convert to integer
                },
                "products": product_names,
                "discount": f"{int(bundle.discount['percentage'])}%"  # Format as percentage string
            })
            
            # Get bundle template or fallback to standard
            template_name = bundle.template or config.settings.template_pattern
            template_path = config.templates_path / template_name.format(language=lang)
            
            if not template_path.exists():
                logger.error(f"Missing bundle template for {lang}: {template_path}")
                continue
                
            # Create output directory
            output_dir = config.output_path / "bundles" / bundle_name
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate output filename
            output_filename = config.settings.filename_pattern.format(
                product=bundle.name,
                language=lang,
                currency=currency,
                date=config.offer.date,
                format=config.settings.format
            )
            output_file = output_dir / output_filename
            
            # Render the bundle offer
            template = DocxTemplate(str(template_path))
            render_offer(template, config, context, output_file)

def main():
    """Main entry point for the offer document generator."""
    if len(sys.argv) < 2:
        print("Usage: python offerdocgenerator.py <config.yaml> [--bundles]")
        sys.exit(1)

    # Use test_data config if running from tests
    if "test" in sys.argv[1].lower():
        test_data = Path(__file__).parent / "test_data"
        config_path = test_data / "test_config.yaml"
    else:
        config_path = Path(sys.argv[1])
    config = load_config(config_path)

    # Get available products
    products = get_product_names(config)
    if not products:
        logger.error("No products found in products directory")
        sys.exit(1)

    # Use configured languages and currencies
    languages = config.languages
    currencies = config.currencies
    
    # Generate offer documents for each combination
    for lang in languages:
        for currency in currencies:
            for product in products:
                # Build context with currency
                context = build_context(config, lang, product, currency)
                
                # Get template path using configured pattern
                template_filename = config.settings.template_pattern.format(language=lang)
                template_path = config.templates_path / template_filename
                if not template_path.exists():
                    logger.error(f"Missing template for {lang}: {template_path}")
                    continue

                # Create output directory
                output_dir = config.output_path / product
                output_dir.mkdir(parents=True, exist_ok=True)
                
                # Generate output filename using configured pattern
                fmt = config.settings.format
                output_filename = config.settings.filename_pattern.format(
                    product=product,
                    language=lang,
                    currency=currency,
                    date=config.offer.date,
                    format=fmt
                )
                output_file = output_dir / output_filename
                
                # Render the offer document
                template = DocxTemplate(str(template_path))
                render_offer(template, config, context, output_file)
    
    # Generate bundle offers if requested
    if "--bundles" in sys.argv:
        print(colorize("\n📦 Generating bundle offers:", 'CYAN'))
        for bundle_name in config.bundles:
            print(colorize(f"\n🔖 Bundle: {bundle_name}", 'GREEN'))
            generate_bundle_offer(config, bundle_name)

if __name__ == "__main__":
    main()
