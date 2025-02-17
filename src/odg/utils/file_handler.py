"""File handling utilities."""

import logging
from pathlib import Path
from typing import Protocol, runtime_checkable, Sequence, Any, Dict, Never, assert_never
import yaml
from docx import Document
from enum import StrEnum, auto

from ..document import (
    DocxDocument,
    DocumentSection,
    DocumentMode,
    SectionType
)

logger = logging.getLogger(__name__)

@runtime_checkable
class ConfigProvider(Protocol):
    """Protocol for configuration providers."""
    base_path: str | Path
    output_dir: str | Path
    folders: Any

class FileError(StrEnum):
    NOT_FOUND = auto()
    PERMISSION_DENIED = auto()
    INVALID_FORMAT = auto()

class FileHandler:
    """Handles file operations for document generation."""
    
    def __init__(self, config: ConfigProvider) -> None:
        """Initialize file handler with configuration."""
        self.config = config
        self.output_dir = Path(config.output_dir)
        self.docx_handler = DocxDocument()
    
    def setup_default_folders(self, base_dir: Path) -> None:
        """Create default folder structure."""
        folders = [
            "templates",
            "common",
            "products",
            "output",
            "generated"
        ]
        
        for folder in folders:
            folder_path = base_dir / folder
            folder_path.mkdir(parents=True, exist_ok=True)
            logger.info(f"Created directory: {folder_path}")

    def generate_config(self, default_config: Dict = None) -> Path:
        """Generate configuration file."""
        config_path = self.output_dir / "config.yaml"
        config_data = default_config or {
            "settings": {
                "base_path": str(self.output_dir),
                "output_prefix": "DOC-",
                "folders": {
                    "templates": "templates",
                    "common": "common",
                    "products": "products",
                    "output": "output",
                    "generated": "generated"
                }
            }
        }
        
        with open(config_path, 'w') as f:
            yaml.dump(config_data, f)
        
        return config_path

    def create_docx_template(self, template_name: str) -> Path:
        """Create a DOCX template file with numbered lists."""
        templates_dir = self.output_dir / "templates"
        templates_dir.mkdir(parents=True, exist_ok=True)
        
        template_path = templates_dir / template_name
        doc = Document()
        
        # Add title
        doc.add_heading("Document Template", 0)
        
        # Add some numbered lists for testing
        p1 = doc.add_paragraph("First Level Item", style='List Number')
        p2 = doc.add_paragraph("Second Level Item", style='List Number 2')
        p3 = doc.add_paragraph("Another First Level Item", style='List Number')
        p4 = doc.add_paragraph("Another Second Level Item", style='List Number 2')
        
        # Apply proper formatting
        for p in [p1, p2, p3, p4]:
            if p.style.name == 'List Number':
                p.paragraph_format.left_indent = 457200  # 0.5 inch
                p.paragraph_format.first_line_indent = -457200  # -0.5 inch
            elif p.style.name == 'List Number 2':
                p.paragraph_format.left_indent = 914400  # 1 inch
                p.paragraph_format.first_line_indent = -457200  # -0.5 inch
            
            font = p.style.font
            font.size = 280000  # 14pt
            
        # Save document
        doc.save(str(template_path))
        return template_path
    
    def create_sample_docx(
        self,
        product_name: str,
        language: str,
        currency: str,
        product_sections: list[DocumentSection] = None,
        product_path: Path = None
    ) -> Path:
        """Create a sample DOCX file with numbered lists."""
        output_dir = self.output_dir / "output"
        output_dir.mkdir(parents=True, exist_ok=True)
        
        output_path = output_dir / f"{product_name}_{language}_{currency}.docx"
        doc = Document()
        
        # Add basic content
        doc.add_heading(product_name, 0)
        doc.add_paragraph(f"Language: {language}")
        doc.add_paragraph(f"Currency: {currency}")
        
        # Add numbered lists
        p1 = doc.add_paragraph("First Level Item", style='List Number')
        p2 = doc.add_paragraph("Second Level Item", style='List Number 2')
        p3 = doc.add_paragraph("Another First Level Item", style='List Number')
        p4 = doc.add_paragraph("Another Second Level Item", style='List Number 2')
        
        # Apply proper formatting
        for p in [p1, p2, p3, p4]:
            if p.style.name == 'List Number':
                p.paragraph_format.left_indent = 457200  # 0.5 inch
                p.paragraph_format.first_line_indent = -457200  # -0.5 inch
            elif p.style.name == 'List Number 2':
                p.paragraph_format.left_indent = 914400  # 1 inch
                p.paragraph_format.first_line_indent = -457200  # -0.5 inch
            
            font = p.style.font
            font.size = 280000  # 14pt
        
        # Save document
        doc.save(str(output_path))
        return output_path
    
    def merge_docx_files(self, files: Sequence[Path], output_path: Path) -> None:
        """Merge multiple DOCX files into one."""
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        merged_doc = Document()
        for file_path in files:
            doc = Document(str(file_path))
            for element in doc.element.body:
                merged_doc.element.body.append(element)
        
        merged_doc.save(str(output_path))

    def validate_config(self, config_path: Path) -> bool:
        """Validate configuration file."""
        try:
            with open(config_path) as f:
                config = yaml.safe_load(f)
            
            # Check required top-level sections
            required_sections = ["settings", "offer", "internationalization"]
            if not all(section in config for section in required_sections):
                return False
                
            # Check settings section
            settings = config.get("settings", {})
            required_settings = ["base_path", "output_prefix", "folders"]
            if not all(setting in settings for setting in required_settings):
                return False
                
            # Check folders configuration
            folders = settings.get("folders", {})
            required_folders = ["templates", "common", "products", "output", "generated"]
            if not all(folder in folders for folder in required_folders):
                return False
                
            return True
            
        except Exception as e:
            logger.error(f"Config validation failed: {e}")
            return False

    def validate_doc_structure(self, doc_path: Path) -> bool:
        """Validate document structure."""
        try:
            doc = Document(str(doc_path))
            return len(doc.paragraphs) > 0
        except Exception as e:
            logger.error(f"Document validation failed: {e}")
            return False

    def has_numbered_lists(self, doc_path: Path) -> bool:
        """Check if document contains numbered lists."""
        try:
            doc = Document(str(doc_path))
            return any(p.style.name.startswith('List Number') for p in doc.paragraphs)
        except Exception as e:
            logger.error(f"Failed to check numbered lists: {e}")
            return False

    def validate_merged_doc(self, doc_path: Path) -> bool:
        """Validate merged document."""
        return self.validate_doc_structure(doc_path)

    def has_required_sections(self, doc_path: Path) -> bool:
        """Check if document has required sections."""
        try:
            doc = Document(str(doc_path))
            return len(doc.sections) > 0
        except Exception as e:
            logger.error(f"Failed to check sections: {e}")
            return False

    def handle_error(self, error: FileError) -> Never:
        match error:
            case FileError.NOT_FOUND:
                raise FileNotFoundError("File not found")
            case FileError.PERMISSION_DENIED:
                raise PermissionError("Permission denied")
            case FileError.INVALID_FORMAT:
                raise ValueError("Invalid file format")
            case _:
                assert_never(error)

    def validate_file(self, path: Path) -> bool:
        if not path.exists():
            self.handle_error(FileError.NOT_FOUND)
        if not os.access(path, os.R_OK):
            self.handle_error(FileError.PERMISSION_DENIED)
        return True
