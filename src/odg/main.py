from __future__ import annotations

import logging
import sys
import argparse
from enum import StrEnum, auto
from pathlib import Path
from typing import Final, NoReturn, List, Optional, TypedDict, NotRequired, Literal, Never, assert_never, LiteralString
from dataclasses import dataclass, field

import yaml
from jinja2 import Environment, FileSystemLoader
from docx import Document

from config.settings import Config
from odg.document_generator import DocumentGenerator
from odg.document import DocumentMode, DocumentSection, SectionType
from odg.utils.file_handler import FileHandler

logger = logging.getLogger(__name__)

class CommandStatus(StrEnum):
    SUCCESS = auto()
    CONFIG_ERROR = auto()
    GENERATION_ERROR = auto()
    VALIDATION_ERROR = auto()

@dataclass(frozen=True, slots=True)
class GenerationOptions:
    language: str
    currency: str
    mode: DocumentMode = field(default=DocumentMode.DRAFT)
    validate: bool = field(default=True)
    sections: List[DocumentSection] = field(default_factory=list)

class DefaultConfig(TypedDict):
    settings: dict[str, dict[str, str] | str]
    offer: dict[str, NotRequired[str | dict[str, str]]]
    customer: dict[str, dict[str, str]]
    sales: dict[str, str]
    internationalization: dict[str, str | list[str]]

class ConfigGenerator:
    VERSION: Final[str] = "1.0.0"

    def __init__(self, output_dir: str | Path) -> None:
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        self.config = Config(
            base_path=str(self.output_dir),
            output_dir=str(self.output_dir)
        )
        self.file_handler = FileHandler(self.config)
        self._setup_jinja()

    def _setup_jinja(self) -> None:
        """Setup Jinja2 environment with proper configuration."""
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.output_dir)),
            autoescape=True,
            trim_blocks=True,
            lstrip_blocks=True,
            keep_trailing_newline=True
        )

    def generate_config(self) -> Path:
        """Generate configuration files and folder structure."""
        logger.info(f"Generating config in {self.output_dir}")
        
        try:
            # Create necessary directories first
            self.file_handler.setup_default_folders(self.output_dir)
            
            # Generate config file
            config_path = self.file_handler.generate_config(self._get_default_config())
            
            # Create base templates
            self.create_base_templates()
            
            return config_path
        except Exception as e:
            logger.error(f"Failed to generate config: {e}")
            raise ConfigGenerationError(f"Config generation failed: {e}") from e

    def create_base_templates(self) -> None:
        """Create base template files."""
        templates = [
            "base_template_en.docx",
            "base_template_de.docx",
            "base_template_fr.docx"
        ]
        
        for template_name in templates:
            try:
                template_path = self.file_handler.create_docx_template(template_name)
                logger.info(f"Created template file: {template_path}")
            except Exception as e:
                logger.error(f"Failed to create template {template_name}: {e}")
                raise ConfigGenerationError(f"Template creation failed: {e}") from e

    def create_docx_template(self, template_name: str) -> Path:
        """Create a DOCX template file with default structure."""
        try:
            doc_path = self.file_handler.create_docx_template(template_name)
            
            # Ensure proper numbered list styles are applied
            doc = Document(str(doc_path))
            
            # Apply proper list numbering styles
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith('List Number'):
                    # Ensure the paragraph has proper list formatting
                    paragraph.style.font.size = 280000  # 14pt
                    if 'List Number 2' in paragraph.style.name:
                        paragraph.paragraph_format.left_indent = 914400  # 1 inch
                        paragraph.paragraph_format.first_line_indent = -457200  # -0.5 inch
                    else:
                        paragraph.paragraph_format.left_indent = 457200  # 0.5 inch
                        paragraph.paragraph_format.first_line_indent = -457200  # -0.5 inch

            doc.save(str(doc_path))
            return doc_path

        except Exception as e:
            logger.error(f"Failed to create DOCX template {template_name}: {e}")
            raise ConfigGenerationError(f"DOCX template creation failed: {e}") from e

    def validate_config(self, config_path: Path) -> bool:
        """Validate configuration file."""
        try:
            return self.file_handler.validate_config(config_path)
        except Exception as e:
            logger.error(f"Config validation failed: {e}")
            return False

    def create_sample_docx(
        self, 
        product_name: str, 
        language: str, 
        currency: str,
        product_sections: Optional[List[DocumentSection]] = None,
        product_path: Optional[Path] = None
    ) -> Path:
        """Create a sample DOCX file with structured content."""
        try:
            # Prepare template variables
            template_vars = {
                "product_name": product_name,
                "language": language,
                "currency": currency,
                "project_duration": "12 months",
                "project_start_date": "2024-01-01",
                "test_environment": "Staging",
                "access_requirements": "VPN Access",
                "audit_locations": "Main Data Center",
                "network_size": "500 nodes",
                "compliance_standards": "ISO 27001"
            }

            # Create document with rendered variables
            doc_path = self.file_handler.create_sample_docx(
                product_name=product_name,
                language=language,
                currency=currency,
                product_sections=product_sections,
                product_path=product_path
            )

            # Render Jinja2 variables in the document
            self._render_document_variables(doc_path, template_vars)

            return doc_path

        except Exception as e:
            logger.error(f"Failed to create sample DOCX: {e}")
            raise ConfigGenerationError(f"Sample DOCX creation failed: {e}") from e

    def _render_document_variables(self, doc_path: Path, variables: dict) -> None:
        """Render Jinja2 variables in the document."""
        try:
            doc = Document(str(doc_path))
            
            # Render variables in paragraphs
            for paragraph in doc.paragraphs:
                if "{{" in paragraph.text and "}}" in paragraph.text:
                    template = self.jinja_env.from_string(paragraph.text)
                    rendered_text = template.render(**variables)
                    paragraph.text = rendered_text

            # Save the document with rendered variables
            doc.save(str(doc_path))

        except Exception as e:
            logger.error(f"Failed to render document variables: {e}")
            raise ConfigGenerationError(f"Variable rendering failed: {e}") from e

    def merge_docx_files(self, files: List[Path], output_path: Path) -> None:
        """Merge multiple DOCX files into one."""
        try:
            self.file_handler.merge_docx_files(files, output_path)
        except Exception as e:
            logger.error(f"Failed to merge DOCX files: {e}")
            raise ConfigGenerationError(f"DOCX merge failed: {e}") from e

    def merge_documents(self, docs: list[Path], output_path: Path) -> Path:
        """Merge multiple documents into one."""
        try:
            self.merge_docx_files(docs, output_path)
            return output_path
        except Exception as e:
            logger.error(f"Document merge failed: {e}")
            raise ConfigGenerationError(f"Document merge failed: {e}") from e

    def validate_product_template(self, template_path: Path) -> bool:
        """Validate product template structure."""
        try:
            return self.file_handler.validate_doc_structure(template_path)
        except Exception as e:
            logger.error(f"Product template validation failed: {e}")
            return False

    def _get_default_config(self) -> DefaultConfig:
        """Get default configuration dictionary."""
        return {
            "settings": {
                "base_path": str(self.output_dir),
                "output_prefix": "DOC-",
                "folders": {
                    "templates": "templates",
                    "common": "common",
                    "products": "products",
                    "output": "output",
                    "generated": "generated"
                }
            },
            "offer": {
                "number": "OFR-{{date}}-{{number}}",
                "date": "{{current_date}}",
                "validity": "30 days",
                "currency": "{{currency}}",
                "language": "{{language}}"
            },
            "customer": {
                "details": {
                    "name": "{{customer_name}}",
                    "address": "{{customer_address}}"
                }
            },
            "sales": {
                "representative": "{{sales_rep}}",
                "email": "{{sales_email}}"
            },
            "internationalization": {
                "default_language": "en",
                "supported_languages": ["en", "de", "fr"]
            }
        }

class OfferDocGenerator:
    VERSION: Final[str] = "1.0.0"

    def __init__(self, config: Config | None = None) -> None:
        self.config = config or Config()
        self.file_handler = FileHandler(self.config)
        self.doc_generator = DocumentGenerator(self.config)

    def setup_project(self, output_dir: str | Path) -> None:
        """Set up a new project structure."""
        output_dir = Path(output_dir)
        logger.info(f"Setting up project in {output_dir}")
        
        try:
            self.file_handler.setup_default_folders(output_dir)
            config_path = self.file_handler.generate_config()
            
            if not self.file_handler.validate_config(config_path):
                raise ConfigValidationError("Config validation failed")

            self.create_base_templates()
            
        except Exception as e:
            logger.error(f"Project setup failed: {e}")
            raise ProjectSetupError(f"Failed to set up project: {e}") from e

def main() -> Never:
    """Main entry point with Never return type."""
    try:
        parser = create_argument_parser()
        args = parser.parse_args()

        match args.command:
            case "create":
                config_generator = ConfigGenerator(args.output_dir)
                config_path = config_generator.generate_config()
                sys.exit(CommandStatus.SUCCESS)
            case "generate":
                project_dir = Path(args.project_dir)
                config = Config(base_path=str(project_dir))
                generator = ConfigGenerator(config)
                sys.exit(CommandStatus.SUCCESS)
            case "validate":
                config_generator = ConfigGenerator(args.directory)
                if not config_generator.validate_config():
                    sys.exit(CommandStatus.VALIDATION_ERROR)
                sys.exit(CommandStatus.SUCCESS)
            case _:
                assert_never(args.command)
    except Exception as e:
        logger.error(f"Error: {e}")
        sys.exit(CommandStatus.GENERATION_ERROR)

# Custom exceptions
class ConfigGenerationError(Exception): pass
class ConfigValidationError(Exception): pass
class ProjectSetupError(Exception): pass

def create_argument_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="OfferDoc Generator")
    subparsers = parser.add_subparsers(dest="command", required=True)

    # Create command
    create_parser = subparsers.add_parser("create", help="Create new project")
    create_parser.add_argument(
        "output_dir",
        help="Output directory for the new project"
    )

    # Generate command
    generate_parser = subparsers.add_parser("generate", help="Generate documents")
    generate_parser.add_argument(
        "project_dir",
        help="Project directory containing config.yaml",
        default=".",
        nargs="?"
    )

    # Validate command
    validate_parser = subparsers.add_parser("validate", help="Validate configuration")
    validate_parser.add_argument(
        "directory",
        help="Project directory containing config.yaml"
    )

    return parser

if __name__ == "__main__":
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    sys.exit(main())
