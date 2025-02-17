import logging
from pathlib import Path
from typing import Optional, List, LiteralString, Never, AsyncIterator, Awaitable
from collections.abc import AsyncIterable
import asyncio
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config.settings import Config, ProductSection
from odg.utils.file_handler import FileHandler

logger = logging.getLogger(__name__)

class DocumentGenerator:
    def __init__(self, config: Config):
        self.config = config

    def validate_template_name(self, name: LiteralString) -> bool:
        """Validate template name using literal string type."""
        if not name.endswith((".docx", ".doc")):
            self.handle_invalid_template(name)
        return True
    
    def handle_invalid_template(self, name: str) -> Never:
        """Handle invalid template with Never return type."""
        raise ValueError(f"Invalid template name: {name}")

    def generate_document(self, product: ProductSection, language: str) -> Path:
        """Generate a document for a specific product and language."""
        try:
            doc = self._create_base_document(product["name"])
            self._add_product_sections(doc, product, language)
            
            output_path = self._get_output_path(product["name"], language)
            doc.save(str(output_path))
            
            logger.info(f"Generated document: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to generate document: {e}")
            raise

    def _create_base_document(self, title: str) -> Document:
        """Create a base document with title."""
        doc = Document()
        
        # Add title
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add styles
        self._add_document_styles(doc)
        
        return doc

    def _add_document_styles(self, doc: Document) -> None:
        """Add custom styles to the document."""
        styles = doc.styles
        
        # Heading 1
        h1_style = styles['Heading 1']
        h1_style.font.size = Pt(16)
        h1_style.font.bold = True
        
        # Heading 2
        h2_style = styles['Heading 2']
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True

    def _add_product_sections(self, doc: Document, product: ProductSection, language: str) -> None:
        """Add all product sections to the document."""
        for section_num, section in enumerate(product["sections"], 1):
            self._add_section(doc, section_num, section, language)

    def _add_section(self, doc: Document, section_num: int, section: str, language: str) -> None:
        """Add a section to the document."""
        heading = doc.add_paragraph(style='Heading 1')
        heading.add_run(f"{section_num}. {section}")
        
        # Add section content
        content = doc.add_paragraph()
        content.add_run(self._get_section_content(section, language))

    def _get_section_content(self, section: str, language: str) -> str:
        """Get localized content for a section."""
        # This would typically load content from templates or a content management system
        return f"Content for {section} in {language}"

    def _get_output_path(self, product_name: str, language: str) -> Path:
        """Generate output file path."""
        filename = (
            f"{self.config.output_prefix}"
            f"{product_name.replace(' ', '_')}"
            f"_{language}.docx"
        )
        return self.config.folders.output / filename

class AsyncDocumentGenerator(DocumentGenerator):
    async def generate_documents(
        self,
        templates: list[str]
    ) -> AsyncIterator[tuple[str, Path]]:
        """Generate documents asynchronously."""
        async with asyncio.TaskGroup() as tg:
            tasks = [
                tg.create_task(self._generate_single(template))
                for template in templates
            ]
        
        for task in tasks:
            name, path = await task
            yield name, path

    async def _generate_single(self, template: str) -> tuple[str, Path]:
        """Generate a single document."""
        async with self._get_template_lock(template):
            return await self._process_template(template)

    @staticmethod
    async def process_batch(
        docs: AsyncIterable[tuple[str, Path]]
    ) -> list[Path]:
        """Process a batch of documents."""
        results = []
        async for name, path in docs:
            results.append(path)
        return results
