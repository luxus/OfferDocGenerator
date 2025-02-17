"""DOCX document handling implementation."""

from pathlib import Path
from typing import Sequence, Any, cast
import logging
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from .base import BaseDocument, DocumentHandler

logger = logging.getLogger(__name__)

class DocxHandler(DocumentHandler):
    """DOCX document handler implementation."""
    
    def __init__(self, doc: Document) -> None:
        self._doc = doc
    
    def save(self, path: Path) -> None:
        """Save document to path."""
        self._doc.save(str(path))
    
    def add_heading(self, text: str, level: int = 1) -> Any:
        """Add heading to document."""
        heading = self._doc.add_heading(text, level)
        if level == 0:  # Title
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return heading
    
    def add_paragraph(self, text: str | None = None) -> Any:
        """Add paragraph to document."""
        return self._doc.add_paragraph(text)

class DocxDocument(BaseDocument):
    """DOCX document implementation."""
    
    def create(self) -> DocumentHandler:
        """Create a new DOCX document."""
        return DocxHandler(Document())
    
    def open(self, path: Path) -> DocumentHandler:
        """Open an existing DOCX document."""
        return DocxHandler(Document(str(path)))
    
    def merge(self, documents: Sequence[Path], output_path: Path) -> None:
        """Merge multiple DOCX documents into one."""
        if not documents:
            raise ValueError("No documents provided for merging")
        
        try:
            merged_doc = Document(str(documents[0]))
            for doc_path in documents[1:]:
                doc = Document(str(doc_path))
                for element in doc.element.body:
                    merged_doc.element.body.append(element)
            
            # Ensure consistent styling
            self._normalize_styles(merged_doc)
            merged_doc.save(str(output_path))
            
        except Exception as e:
            logger.error(f"Failed to merge documents: {e}")
            raise
    
    def validate(self, path: Path) -> bool:
        """Validate DOCX document structure."""
        try:
            doc = Document(str(path))
            return (
                self._has_valid_structure(doc) and
                self._has_consistent_numbering(doc) and
                self._has_required_sections(doc)
            )
        except Exception as e:
            logger.error(f"Document validation failed: {e}")
            return False
    
    def _normalize_styles(self, doc: Document) -> None:
        """Ensure consistent styling throughout the document."""
        # Normalize heading styles
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                font = paragraph.style.font
                font.size = Pt(14)
                font.name = 'Arial'
    
    def _has_valid_structure(self, doc: Document) -> bool:
        """Check if document has valid structure."""
        return any(
            p.style.name.startswith('Heading')
            for p in doc.paragraphs
        )
    
    def _has_consistent_numbering(self, doc: Document) -> bool:
        """Check if document has consistent numbering."""
        numbered_lists = [
            p for p in doc.paragraphs
            if p.style.name.startswith('List Number')
        ]
        return bool(numbered_lists)
    
    def _has_required_sections(self, doc: Document) -> bool:
        """Check if document has all required sections."""
        required = {'Introduction', 'Overview', 'Technical Specifications'}
        headings = {
            p.text.strip()
            for p in doc.paragraphs
            if p.style.name.startswith('Heading')
        }
        return required.issubset(headings)